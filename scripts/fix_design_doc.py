#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""直接修改软件详细设计书.docx中与项目实际逻辑不符的内容"""

from docx import Document
from docx.shared import Pt

doc = Document(r'D:\Second-hand Transaction\软件详细设计书.docx')

def fix_para(para, old, new):
    """替换段落文本，保持第一个run的格式"""
    if old not in para.text:
        return False
    full = para.text
    # 清空所有run
    for run in para.runs:
        run.text = ''
    # 将新文本写入第一个run
    if para.runs:
        para.runs[0].text = full.replace(old, new)
    else:
        run = para.add_run(full.replace(old, new))
    return True

# ========== 修改清单 ==========
fixes = []

# 1. MySQL版本
for i, p in enumerate(doc.paragraphs):
    if 'MySQL 8.x' in p.text:
        fix_para(p, 'MySQL 8.x', 'MySQL 9.7')
        print(f'[P{i}] MySQL 8.x → MySQL 9.7')
        break

# 2. exe安装 → 脚本启动
for i, p in enumerate(doc.paragraphs):
    if 'exe安装包实现一键安装' in p.text:
        fix_para(p, '支持通过exe安装包实现一键安装', '支持通过PowerShell脚本（start.ps1）一键启动前后端服务')
        print(f'[P{i}] exe安装 → 脚本启动')
        break

# 3. 删除Flyway定义(项目使用JPA ddl-auto而非Flyway)
for i, p in enumerate(doc.paragraphs):
    if 'Flyway' in p.text and '数据库版本迁移' in p.text:
        fix_para(p, 'Flyway：数据库版本迁移工具。',
                 'BCrypt：密码哈希加密算法，用于用户密码安全存储。')
        print(f'[P{i}] Flyway → BCrypt')
        break

# 4. Windows exe安装程序 → Vite前端开发服务器
for i, p in enumerate(doc.paragraphs):
    if 'Windows exe 安装程序' in p.text:
        fix_para(p, '（5）Windows exe 安装程序', '（5）Vite 前端开发服务器')
        print(f'[P{i}] exe安装程序 → Vite')
        break

# 5. 物流模块描述修正
for i, p in enumerate(doc.paragraphs):
    if '物流模块负责模拟物流轨迹返回' in p.text:
        fix_para(p, '物流模块负责模拟物流轨迹返回；', '物流模块负责运单管理和物流轨迹查询（Mock模拟）；')
        print(f'[P{i}] 物流描述修正')
        break

# 6. 启动模块描述修正
for i, p in enumerate(doc.paragraphs):
    if '启动模块负责自动化启动服务' in p.text:
        fix_para(p, '启动模块负责自动化启动服务。', '管理后台模块负责提供管理员后台管理功能。')
        print(f'[P{i}] 启动模块描述修正')
        break

# 7. (6)系统启动模块 → (6)管理后台模块
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '（6）系统启动模块':
        fix_para(p, '（6）系统启动模块', '（6）管理后台模块')
        print(f'[P{i}] 子系统6名称修正')
        break

for i, p in enumerate(doc.paragraphs):
    if '负责系统启动、环境检查、启动后端服务、启动前端页面、自动创建管理员' in p.text:
        fix_para(p, '负责系统启动、环境检查、启动后端服务、启动前端页面、自动创建管理员',
                 '负责管理员仪表盘、用户管理、商品管理、订单管理、售后处理和举报审核等功能。')
        print(f'[P{i}] 子系统6描述修正')
        break

# 8. 3.6 节标题
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '3.6 系统启动模块':
        fix_para(p, '3.6 系统启动模块', '3.6 管理后台模块')
        print(f'[P{i}] 3.6标题修正')
        break

# 9. 3.6 模块名称
for i, p in enumerate(doc.paragraphs):
    if '模块名称：系统启动模块' in p.text:
        fix_para(p, '模块名称：系统启动模块', '模块名称：管理后台模块')
        print(f'[P{i}] M6名称修正')
        break

# 10. 3.6 输入
for i, p in enumerate(doc.paragraphs):
    if '启动命令、运行环境参数。' in p.text and i > 200:
        fix_para(p, '启动命令、运行环境参数。', '管理操作指令、筛选条件、审核参数。')
        print(f'[P{i}] M6输入修正')
        break

# 11. 3.6 处理的4个子项 - (1)到(4)
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '（1）检查 MySQL 是否运行。':
        fix_para(p, '（1）检查 MySQL 是否运行。', '（1）仪表盘数据统计：汇总用户数、商品数、订单数等。')
        print(f'[P{i}] M6处理1')
    if p.text.strip() == '（2）启动 Spring Boot 后端。':
        fix_para(p, '（2）启动 Spring Boot 后端。', '（2）用户管理：查看、搜索、禁用/启用用户账号。')
        print(f'[P{i}] M6处理2')
    if p.text.strip() == '（3）启动前端页面。':
        fix_para(p, '（3）启动前端页面。', '（3）商品管理：查看、搜索、下架/上架商品。')
        print(f'[P{i}] M6处理3')
    if p.text.strip() == '（4）自动创建管理员。':
        fix_para(p, '（4）自动创建管理员。', '（4）订单与售后管理：查看订单状态，处理售后申请。')
        print(f'[P{i}] M6处理4')

# 12. 3.6 算法描述
for i, p in enumerate(doc.paragraphs):
    if '启动算法' in p.text and '算法描述' in doc.paragraphs[i-1].text if i>0 else False:
        fix_para(p, '启动算法：', '权限校验算法：')
        print(f'[P{i}] 算法名修正')

for i, p in enumerate(doc.paragraphs):
    if '依次检查端口占用' in p.text:
        fix_para(p, '依次检查端口占用；', '校验当前用户角色是否为ADMIN；')
        print(f'[P{i}] 启动算法步骤1修正')
        break
for p in doc.paragraphs:
    if '编译并打包 Jar' in p.text:
        fix_para(p, '编译并打包 Jar；', '根据角色返回对应管理功能和数据；')
        print(f'算法步骤2修正')
        break
for p in doc.paragraphs:
    if '等待数据库启动完成后继续启动' in p.text:
        fix_para(p, '等待数据库启动完成后继续启动。', '记录管理员操作日志以供审计。')
        print(f'算法步骤3修正')
        break

# 13. 3.6 输出
for i, p in enumerate(doc.paragraphs):
    if '系统启动状态、日志信息' in p.text:
        fix_para(p, '系统启动状态、日志信息。', '管理操作结果、统计数据与处理状态。')
        print(f'[P{i}] M6输出修正')
        break

# 14. API响应格式 4.2节 - {code, message, data} → {success, data/error}
for i, p in enumerate(doc.paragraphs):
    if '接口统一采用 ApiResponse 格式' in p.text:
        fix_para(p, '接口统一采用 ApiResponse 格式。',
                 '接口统一采用 ApiResponse 格式：{ success: true, data: ... } 或 { success: false, error: { code, message } }。')
        print(f'[P{i}] API响应格式1修正')
        break

# 修改下面紧跟的 { code, message, data }  三行
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == 'ApiResponse 结构：':
        # 修改后面三个段落
        for j in range(i+1, min(i+5, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip()
            if txt == 'code,':
                fix_para(doc.paragraphs[j], 'code,', '  "success": true,')
            elif txt == 'message,':
                fix_para(doc.paragraphs[j], 'message,', '  "data": { ... }')
            elif txt == 'data':
                fix_para(doc.paragraphs[j], 'data', '// 或失败时: { "success": false, "error": { "code": "...", "message": "..." } }')
        print(f'[P{i}] API响应格式2修正')
        break

# 15. 6.1节错误格式
for i, p in enumerate(doc.paragraphs):
    if '统一错误返回格式' in p.text:
        # 修改后面两行
        for j in range(i+1, min(i+4, len(doc.paragraphs))):
            txt = doc.paragraphs[j].text.strip()
            if txt == 'code,':
                fix_para(doc.paragraphs[j], 'code,', '  "success": false,')
            elif txt == 'message':
                fix_para(doc.paragraphs[j], 'message', '  "error": { "code": "...", "message": "..." }')
        print(f'[P{i}] 错误格式修正')
        break

# 16. 数据库字段名
for i, p in enumerate(doc.paragraphs):
    if 'cover_url' in p.text:
        fix_para(p, 'cover_url', 'cover_image_url')
        print(f'[P{i}] cover_url → cover_image_url')
for i, p in enumerate(doc.paragraphs):
    if 'company_code' in p.text:
        fix_para(p, 'company_code', 'carrier_code')
        print(f'[P{i}] company_code → carrier_code')

# 17. 商品输入字段
for i, p in enumerate(doc.paragraphs):
    if '商品标题、商品价格、商品描述、商品封面图 URL' in p.text:
        fix_para(p, '商品标题、商品价格、商品描述、商品封面图 URL',
                 '商品标题、价格、描述、封面图、分类、成色、数量、免邮/邮费')
        print(f'[P{i}] 商品输入字段修正')
        break

# 18. 测试计划表修正
if doc.tables:
    t = doc.tables[0]
    mappings = {
        '火车票模块': ('商品管理模块', '验证商品发布、编辑、列表查询、详情查看等功能'),
        '酒店预订模块': ('订单管理模块', '验证订单创建、支付、发货、状态流转等功能'),
        '火车餐饮模块': ('售后管理模块', '验证售后申请、审批、拒绝、状态跟踪等功能'),
        '消息通知模块': ('管理后台模块', '验证仪表盘、用户/商品/订单/售后管理等功能'),
    }
    for row in t.rows:
        cell_text = row.cells[2].text.strip()
        if cell_text in mappings:
            new_name, new_desc = mappings[cell_text]
            row.cells[2].text = new_name
            row.cells[3].text = new_desc
            print(f'测试表: {cell_text} → {new_name}')

# ===== 保存 =====
doc.save(r'D:\Second-hand Transaction\软件详细设计书.docx')
print('\n===== 修改完成！=====')
