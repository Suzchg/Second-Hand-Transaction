#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""生成二手交易平台需求规格报告.docx"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import datetime

doc = Document()

# ===== 全局样式设置 =====
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '宋体')

# 设置标题样式
for i in range(1, 5):
    heading = doc.styles[f'Heading {i}']
    hfont = heading.font
    hfont.name = '黑体'
    hfont.color.rgb = RGBColor(0, 0, 0)
    heading.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')
    if i == 1:
        hfont.size = Pt(18)
    elif i == 2:
        hfont.size = Pt(15)
    elif i == 3:
        hfont.size = Pt(13)
    elif i == 4:
        hfont.size = Pt(12)

def heading(text, level=1):
    """添加标题并设置中文样式"""
    h = doc.add_heading(text, level=level)
    # 确保每个标题都正确设置字体
    for run in h.runs:
        run.font.name = '黑体'
        run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')
    return h

def para(text, bold=False):
    """添加正文段落"""
    p = doc.add_paragraph(text)
    if bold:
        for run in p.runs:
            run.bold = True
    return p

def add_table(headers, rows):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers), style='Table Grid')
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    # 数据行
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            table.rows[r + 1].cells[c].text = str(val)
    doc.add_paragraph()
    return table

# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('二手交易平台')
run.font.size = Pt(28)
run.font.name = '黑体'
run.bold = True
run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('软件需求规格说明书')
run.font.size = Pt(22)
run.font.name = '黑体'
run.element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', '黑体')

doc.add_paragraph()
doc.add_paragraph()

# 版本信息
info_lines = [
    f'版本号：V1.0',
    f'日期：{datetime.date.today().strftime("%Y年%m月%d日")}',
    '状态：已完成',
]
for line in info_lines:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    run.font.size = Pt(14)

doc.add_page_break()

# ============================================================
# 修订历史
# ============================================================
heading('修订记录', level=1)

add_table(
    ['版本', '日期', '修订人', '修订说明'],
    [
        ['V1.0', datetime.date.today().strftime('%Y-%m-%d'), '系统开发组', '初始版本，覆盖全部已实现功能'],
    ]
)

doc.add_page_break()

# ============================================================
# 第 1 章：范围
# ============================================================
heading('1 范围', level=1)

heading('1.1 标识', level=2)
para('项目名称：二手交易平台（SecondHand）')
para('项目类型：Web 在线交易系统')
para('开发语言：Java（后端）+ JavaScript（前端）')
para('英文标识：secondhand-backend / secondhand-frontend')

heading('1.2 系统概述', level=2)
para(
    '二手交易平台是一个基于 Web 的 C2C 二手商品在线交易系统。系统采用前后端分离架构：'
    '后端基于 Spring Boot 3.3 框架提供 RESTful API，使用 Spring Security + JWT 实现身份认证与授权；'
    '前端基于 Vue 3 + Vite 框架构建单页面应用（SPA），提供响应式用户界面。'
    '数据库采用 MySQL 关系型数据库存储业务数据。'
)
para(
    '系统核心业务涵盖用户注册登录、商品发布浏览、订单创建与支付、物流追踪、售后退款退货等完整交易闭环。'
    '同时提供了评论互动、收藏关注、私聊沟通、出价议价、评分评价、举报管理等社区化功能，'
    '以及管理员后台进行用户管理、商品审核、订单监控、售后处理和举报处置。'
)

heading('1.3 文档概述', level=2)
para(
    '本文档为"二手交易平台"的软件需求规格说明书（Software Requirements Specification, SRS），'
    '依据国家软件工程标准规范编制。文档内容包括系统的功能需求、性能需求、接口需求、数据需求、'
    '安全需求、运行环境需求以及软件质量因素等内容，旨在为后续设计、开发和测试阶段提供完整的需求基线。'
)
para('本文档的预期读者包括：')
para('（1）项目指导教师 — 用于评估系统功能完整性和文档规范性；')
para('（2）开发人员 — 作为系统设计、编码和测试的依据；')
para('（3）测试人员 — 作为编写测试用例和验收测试的参考标准。')

heading('1.4 术语', level=2)
add_table(
    ['术语/缩写', '全称', '说明'],
    [
        ['SRS', 'Software Requirements Specification', '软件需求规格说明书'],
        ['CSCI', 'Computer Software Configuration Item', '计算机软件配置项'],
        ['SPA', 'Single Page Application', '单页面应用程序'],
        ['JWT', 'JSON Web Token', 'JSON Web 令牌，用于无状态身份认证'],
        ['REST', 'Representational State Transfer', '表述性状态转移，API 设计风格'],
        ['ORM', 'Object-Relational Mapping', '对象关系映射'],
        ['JPA', 'Java Persistence API', 'Java 持久化 API'],
        ['BCrypt', '—', '密码哈希加密算法'],
        ['C2C', 'Consumer to Consumer', '个人对个人的交易模式'],
    ]
)

doc.add_page_break()

# ============================================================
# 第 2 章：引用文件
# ============================================================
heading('2 引用文件', level=1)
para('（1）GB/T 9385-2008 计算机软件需求规格说明规范')
para('（2）GB/T 8567-2006 计算机软件文档编制规范')
para('（3）IEEE Std 830-1998 软件需求规格说明书推荐实践')
para('（4）Spring Boot 3.3.2 官方文档')
para('（5）Vue.js 3.5 官方文档')
para('（6）MySQL 9.0 官方参考手册')

doc.add_page_break()

# ============================================================
# 第 3 章：需求
# ============================================================
heading('3 需求', level=1)

heading('3.1 需求状态和方式', level=2)
para(
    '本文档中的所有需求均已通过代码实现验证。需求来源于二手交易业务的通用场景分析，'
    '结合电子商务平台的行业标准和用户操作习惯进行定义。各项功能需求均有对应的后端 API 端点和前端页面实现。'
    '需求优先级定义如下：'
)
para('• 关键需求（P0）：系统核心业务流程必需，不可缺失。如用户认证、商品发布、订单创建。')
para('• 重要需求（P1）：提升用户体验的必要功能。如收藏、评论、物流追踪。')
para('• 增强需求（P2）：锦上添花的功能。如私聊、评分、出价议价。')

heading('3.2 需求概述', level=2)

heading('3.2.1 目标', level=3)
para(
    '本系统旨在为个人用户提供一个安全、便捷的二手商品在线交易平台。'
    '卖家和买家可以通过平台完成商品的发布、浏览、下单、支付、发货、收货以及售后退款退货等完整交易流程。'
    '同时，系统提供商品评论、收藏、私聊沟通、出价议价、用户评分、举报管理等社区化功能，'
    '构建一个可信赖的二手交易社区生态。'
)
para('系统主要目标包括：')
para('（1）实现完整的 C2C 二手商品交易闭环，涵盖商品发布到售后处理的全流程；')
para('（2）提供安全可靠的用户认证机制，保护用户账户和数据安全；')
para('（3）支持基于商品分类、关键词、价格区间的灵活商品检索；')
para('（4）保障交易双方的合法权益，通过售后机制和举报机制维护平台秩序；')
para('（5）提供管理员后台，实现对用户、商品、订单、举报的全方位管理和监控；')
para('（6）支持移动端和桌面端的响应式访问，适配不同屏幕尺寸。')

heading('3.2.2 运行环境', level=3)
para('本系统采用 B/S（浏览器/服务器）架构，用户通过浏览器访问前端页面，前端通过 HTTP 协议调用后端 API。', bold=True)

para('【硬件环境】')
para('服务器端：')
para('• CPU：支持 x86_64 架构的现代处理器')
para('• 内存：最低 2GB，推荐 4GB 及以上')
para('• 存储：最低 1GB 可用空间（不含数据库存储）')
para('客户端：')
para('• 任何支持现代浏览器的设备（PC、平板、手机）')
para('• 屏幕分辨率：最低 375px 宽度（移动端适配）')

para('【软件环境】')
para('服务器端：')
add_table(
    ['软件', '版本/说明'],
    [
        ['操作系统', 'Windows 10/11、Linux、macOS'],
        ['Java 运行环境', 'JDK 17 及以上（当前使用 JDK 24）'],
        ['数据库', 'MySQL 9.7（兼容 8.0+）'],
        ['Web 服务器', '内嵌 Tomcat 10（Spring Boot 自带）'],
        ['构建工具', 'Maven 3.9+（使用 Maven Wrapper）'],
    ]
)

para('客户端：')
add_table(
    ['软件', '版本/说明'],
    [
        ['浏览器', 'Chrome 90+、Edge 90+、Firefox 88+、Safari 14+'],
        ['JavaScript', '需启用 JavaScript 和 Cookies'],
    ]
)

heading('3.2.3 用户特点', level=3)
para('系统面向以下用户群体：')
para('（1）普通买家用户：需要通过平台浏览和购买二手商品的个人用户。'
    '这类用户需要直观的商品浏览界面、便捷的下单流程和安全的支付体验。'
    '用户技术水平参差不齐，界面设计需简洁直观、易于操作。')
para('（2）卖家用户：需要在平台上发布和管理商品出售信息的个人用户。'
    '这类用户需要方便的商品发布表单、订单管理工具和收入查看功能。')
para('（3）系统管理员：负责平台日常运营和管理的后台用户。'
    '需要在后台管理界面中进行用户管理、商品审核、订单监控、售后处理和举报处置等操作。'
    '管理员需要具备一定的计算机操作基础。')

heading('3.2.4 关键点', level=3)
para('系统关键功能模块如下：')

para('（一）用户认证模块', bold=True)
para('• 用户注册：支持手机号/邮箱注册，密码 BCrypt 加密存储')
para('• 用户登录：JWT 令牌认证，支持令牌过期自动刷新')
para('• 密码修改：验证旧密码后修改为新密码')
para('• 角色管理：区分普通用户（USER）和管理员（ADMIN）角色')

para('（二）商品管理模块', bold=True)
para('• 商品发布：卖家填写商品标题、描述、价格、成色、分类、图片等信息')
para('• 商品浏览：分页展示在售商品列表，支持按分类筛选和关键词搜索')
para('• 商品详情：展示商品全部信息，包括多图浏览（ImageGallery 组件）和卖家信息')
para('• 商品编辑：卖家可修改自己发布的商品信息')
para('• 商品状态管理：支持草稿、在售、下架三种状态')
para('• 商品分类：多级商品分类体系，便于商品组织和检索')
para('• 商品图片：支持多图上传和图片缩放功能')

para('（三）订单管理模块', bold=True)
para('• 订单创建：买家选择商品后填写收货信息下单')
para('• 支付占位：模拟支付操作，将订单状态从"已创建"变更为"已支付"')
para('• 卖家发货：卖家录入快递公司和运单号完成发货')
para('• 确认收货：买家确认收到商品完成交易')
para('• 订单取消：在特定状态下允许取消订单')

para('（四）售后管理模块', bold=True)
para('• 售后申请：买家在订单完成后可申请退款或退货退款')
para('• 售后处理：卖家可审批通过或拒绝售后申请')
para('• 售后状态跟踪：记录申请、审批、关闭全流程状态')

para('（五）物流追踪模块', bold=True)
para('• 运单创建：卖家发货时自动创建物流运单')
para('• 物流查询：根据订单查询物流轨迹信息（当前使用 Mock 模拟数据）')
para('• 快递公司支持：支持多家快递公司编码')

para('（六）评论互动模块', bold=True)
para('• 商品评论：买家可对购买过的商品发表文字评论')
para('• 评论展示：在商品详情页展示该商品的用户评论列表')

para('（七）收藏关注模块', bold=True)
para('• 商品收藏：用户可收藏感兴趣的商品')
para('• 收藏列表：用户可查看和管理自己的收藏列表')
para('• 取消收藏：支持移除已收藏的商品')

para('（八）私聊通讯模块', bold=True)
para('• 用户私聊：买家可与卖家进行一对一即时通讯')
para('• 消息中心：统一查看和管理所有聊天消息')
para('• 消息列表：按会话组织消息，展示最新消息摘要')

para('（九）出价议价模块', bold=True)
para('• 买家出价：买家可对在售商品提出还价')
para('• 出价状态：支持待处理、接受、拒绝等状态流转')

para('（十）评分评价模块', bold=True)
para('• 用户评分：交易完成后买家可对卖家进行评分')
para('• 评分展示：在卖家主页展示其综合评分和历史评价')

para('（十一）举报管理模块', bold=True)
para('• 用户举报：用户可举报违规商品或不良行为')
para('• 举报处理：管理员在后台审核并处理举报信息')
para('• 举报类型：多种举报原因供选择（虚假信息、违禁品、欺诈行为等）')

para('（十二）管理后台模块', bold=True)
para('• 数据仪表盘：展示平台核心运营数据概览')
para('• 用户管理：查看、搜索、禁用/启用的用户管理功能')
para('• 商品管理：审核、下架/上架所有商品')
para('• 订单管理：查看和监控全部订单状态')
para('• 售后管理：处理平台所有售后申请')
para('• 举报管理：审核和处理用户举报信息')

para('（十三）用户个人信息模块', bold=True)
para('• 个人资料：查看和编辑个人资料（昵称、手机号、邮箱、头像）')
para('• 地址管理：新增、编辑、删除收货地址')
para('• 地区选择：三级联动地区选择器（省/市/区）')

para('（十四）其他功能', bold=True)
para('• 卖家历史商品：查看某卖家发布的所有在售商品')
para('• 切换账号：支持退出当前账号后切换登录')
para('• 售后政策：展示平台售后政策和规则说明')
para('• 隐私政策：展示平台隐私保护政策')

para('【关键技术实现】', bold=True)
para('• Web 开发框架：后端使用 Spring Boot 3.3.2 框架实现 RESTful API；前端使用 Vue 3 框架和 Vite 构建工具实现 SPA')
para('• 安全认证：使用 Spring Security 6.x + JWT (jjwt 0.12.6) 实现无状态身份认证和基于角色的访问控制')
para('• 数据库技术：使用 MySQL 关系型数据库存储用户信息、商品信息、订单数据等，通过 Spring Data JPA + Hibernate 6.5 实现 ORM')
para('• 密码加密：用户密码使用 BCrypt 算法进行哈希加密存储，保障安全性')
para('• 图片管理：商品图片支持多图上传，采用前端 ImageUploader 组件处理图片选择与预览，图片路径存储于数据库中')
para('• 物流查询：采用策略模式（Strategy Pattern）设计物流查询接口，当前使用 Mock 实现模拟物流数据，方便后续接入真实物流 API（如快递 100）')
para('• 分页查询：商品列表、订单列表等采用 Spring Data 分页机制，支持按页和每页条数灵活查询')
para('• API 统一响应：所有 API 均返回统一结构的 JSON 响应体 { success, data, error }')
para('• 全局异常处理：通过 @RestControllerAdvice 实现统一的异常捕获和错误响应格式化')

heading('3.2.5 约束条件', level=3)
para('（1）经费与资源约束：本项目为课程设计项目，无额外经费投入。开发和部署均使用免费开源软件和个人计算机。')
para('（2）开发周期约束：项目需在学期内完成，开发时间有限，因此采用快速原型迭代开发模式。')
para('（3）技术约束：')
para('    • 后端语言必须使用 Java，框架使用 Spring Boot；')
para('    • 前端框架使用 Vue 3；')
para('    • 数据库使用 MySQL 关系型数据库；')
para('    • 所有依赖均为开源免费组件。')
para('（4）法律合规约束：系统需遵守《个人信息保护法》相关规定，保护用户个人信息安全与隐私。')
para('（5）运行约束：系统需支持 Windows 环境运行，前端需兼容主流浏览器。')

heading('3.3 需求规格', level=2)

heading('3.3.1 系统整体功能/结构', level=3)
para('本系统采用前后端分离的 B/S 架构，整体划分为以下几个层次：')
para('（1）表现层（Frontend）：基于 Vue 3 的单页面应用，负责用户界面展示和交互，通过 HTTP 与后端 API 通信。')
para('（2）网关层（Security）：基于 Spring Security Filter Chain 实现请求拦截、JWT 认证和权限控制。')
para('（3）业务逻辑层（Service）：实现核心业务逻辑，包含认证、商品、订单、售后、物流、评论、收藏、私聊、出价、评分、举报等模块。')
para('（4）数据访问层（Repository）：基于 Spring Data JPA 实现数据持久化操作。')
para('（5）数据存储层（Database）：使用 MySQL 关系型数据库存储业务数据。')

para('系统架构图如下：')
para(
    '┌─────────────────────────────────────────────────┐\n'
    '│  浏览器 (Vue 3 SPA)  :5173                       │\n'
    '│  ┌──────┐ ┌──────┐ ┌──────┐ ┌──────┐          │\n'
    '│  │ 首页  │ │ 商品  │ │ 订单  │ │ 消息  │  ...    │\n'
    '│  └──────┘ └──────┘ └──────┘ └──────┘          │\n'
    '└──────────────────┬──────────────────────────────┘\n'
    '                   │ HTTP / JSON\n'
    '┌──────────────────▼──────────────────────────────┐\n'
    '│  Spring Boot Backend  :8088                      │\n'
    '│  ┌─────────────────────────────────────────┐    │\n'
    '│  │  Security Filter Chain + JWT Auth       │    │\n'
    '│  ├─────────────────────────────────────────┤    │\n'
    '│  │  Controller Layer (REST API)            │    │\n'
    '│  ├─────────────────────────────────────────┤    │\n'
    '│  │  Service Layer (Business Logic)         │    │\n'
    '│  ├─────────────────────────────────────────┤    │\n'
    '│  │  Repository Layer (JPA + Hibernate)     │    │\n'
    '│  └─────────────────────────────────────────┘    │\n'
    '└──────────────────┬──────────────────────────────┘\n'
    '                   │ JDBC\n'
    '┌──────────────────▼──────────────────────────────┐\n'
    '│  MySQL Database  :3306                           │\n'
    '│  数据库: secondhand                               │\n'
    '└─────────────────────────────────────────────────┘\n'
)

heading('3.3.2 各子系统功能/结构', level=3)
para('系统按照业务功能划分为以下 14 个子系统模块：')

add_table(
    ['序号', '模块名称', 'Java 包路径', '功能概要'],
    [
        ['1', '用户认证模块', 'auth', '注册、登录、JWT 令牌管理、密码修改、角色认证'],
        ['2', '用户管理模块', 'user', '个人资料、头像管理、收货地址管理、地区数据'],
        ['3', '商品管理模块', 'product', '商品 CRUD、状态管理、分类管理、多图管理'],
        ['4', '订单管理模块', 'order', '订单创建、支付、发货、确认收货、状态流转日志'],
        ['5', '售后管理模块', 'aftersale', '退款/退货退款申请、审批、拒绝、状态跟踪'],
        ['6', '物流追踪模块', 'logistics', '运单管理、物流轨迹查询（策略模式设计）'],
        ['7', '评论互动模块', 'comment', '商品评论的发布与展示'],
        ['8', '收藏关注模块', 'favorite', '商品收藏与取消收藏、收藏列表管理'],
        ['9', '私聊通讯模块', 'chat', '买卖双方一对一私聊、消息中心、会话管理'],
        ['10', '出价议价模块', 'offer', '买家出价、出价状态管理'],
        ['11', '评分评价模块', 'rating', '交易后评分、卖家综合评价展示'],
        ['12', '举报管理模块', 'report', '用户举报提交、管理员举报审核处理'],
        ['13', '管理后台模块', 'admin', '仪表盘、用户/商品/订单/售后/举报综合管理'],
        ['14', '支付模块', 'payment', '支付占位接口（当前为模拟实现）'],
    ]
)

heading('3.3.3 规则约定', level=3)
para('（1）API 设计规范：所有 API 遵循 RESTful 设计风格，统一使用 JSON 格式的请求体和响应体。')
para('（2）统一响应格式：所有接口返回 ApiResponse<T> 结构：')
para('    { "success": true, "data": T }          // 成功')
para('    { "success": false, "error": { "code": "ERROR_CODE", "message": "..." } }  // 失败')
para('（3）认证方式：除注册、登录和公开查询接口外，所有 API 需携带 Authorization: Bearer <jwt_token> 请求头。')
para('（4）错误码规范：系统定义了统一的错误码体系，包括 VALIDATION_ERROR、INVALID_CREDENTIALS、FORBIDDEN、NOT_FOUND、CONFLICT、INTERNAL_ERROR 等。')
para('（5）代码组织：后端遵循 Package-by-Feature 架构，每个业务模块内部按 Controller → Service → Repository → Entity 分层组织。')
para('（6）前端路由：使用 Vue Router 实现 SPA 路由，通过路由守卫（beforeEach）实现未登录拦截和管理员权限控制。')
para('（7）数据库约定：')
para('    • 金额以"分"为单位存储（price_cent, amount_cent），避免浮点精度问题；')
para('    • 状态字段使用 VARCHAR 存储枚举值，便于阅读和扩展；')
para('    • 所有表均包含 created_at 和 updated_at 字段记录时间戳；')
para('    • 使用 ddl-auto: update 策略自动同步实体与表结构。')

heading('3.4 CSCI 需求', level=2)
heading('3.4.1 用户认证模块（auth）', level=3)

para('【功能描述】', bold=True)
para('提供用户注册、登录、JWT 令牌管理和密码修改功能，是整个系统的安全基础。')

para('【输入要求】', bold=True)
para('• 注册：身份类型（PHONE/EMAIL）、标识符（手机号或邮箱）、密码（6-64 位）')
para('• 登录：身份类型、标识符、密码')
para('• 修改密码：旧密码、新密码（6-64 位）')

para('【处理要求】', bold=True)
para('• 注册时密码使用 BCrypt 算法加密存储')
para('• 登录时验证 BCrypt 哈希，成功后颁发 JWT 令牌（有效期 120 分钟）')
para('• JWT 令牌包含用户 ID 和角色信息，签名使用 HMAC-SHA256 算法')
para('• 支持基于角色的访问控制（USER / ADMIN）')
para('• 系统启动时自动创建默认管理员账号（admin / admin123）')

para('【输出要求】', bold=True)
para('• 注册成功：返回 JWT 令牌、用户 ID、昵称')
para('• 登录成功：返回 JWT 令牌、用户 ID、昵称')
para('• 获取当前用户：返回用户 ID、昵称（token 字段为 null）')

para('【性能指标】', bold=True)
para('• 响应时间：≤ 2 秒')
para('• 并发支持：≥ 1000 请求/分钟')

para('【异常处理】', bold=True)
para('• 手机号/邮箱已注册 → 409 IDENTITY_EXISTS')
para('• 账号或密码错误 → 401 INVALID_CREDENTIALS')
para('• 账号已禁用 → 403 FORBIDDEN')
para('• 参数校验失败 → 400 VALIDATION_ERROR')

heading('3.4.2 用户管理模块（user）', level=3)

para('【功能描述】', bold=True)
para('提供用户个人资料管理、头像上传、收货地址管理（CRUD）以及多级地区数据查询功能。')

para('【输入要求】', bold=True)
para('• 个人资料：昵称、手机号、邮箱、头像 URL')
para('• 收货地址：收货人姓名、手机号、省/市/区编码、详细地址')

para('【处理要求】', bold=True)
para('• 用户可查看和编辑自己的个人资料')
para('• 用户可新增、编辑、删除收货地址')
para('• 支持三级联动地区选择（省-市-区）')

para('【输出要求】', bold=True)
para('• 返回用户个人信息和地址列表')

heading('3.4.3 商品管理模块（product）', level=3)

para('【功能描述】', bold=True)
para('提供商品发布、浏览、搜索、编辑、状态管理和多级分类管理功能，是系统的核心业务模块。')

para('【输入要求】', bold=True)
para('• 发布商品：标题（≤100字符）、价格（分）、封面图 URL、分类 ID、描述、成色、库存数量、是否免邮、邮费（分）')
para('• 商品查询：分页参数（page, size）、分类筛选、关键词搜索')

para('【处理要求】', bold=True)
para('• 商品状态管理：草稿（DRAFT）→ 在售（ON_SALE）→ 下架（OFF_SALE）')
para('• 支持商品成色标识：全新、几乎全新、有使用痕迹、有明显瑕疵等')
para('• 支持商品多级分类体系')
para('• 支持多张商品图片上传和管理')
para('• 卖家仅可编辑自己发布的商品')
para('• 管理员可下架/上架任何商品')

para('【输出要求】', bold=True)
para('• 商品列表：分页返回，包含商品基本信息')
para('• 商品详情：完整商品信息，包含图片列表、卖家信息')

para('【性能指标】', bold=True)
para('• 商品列表查询响应时间：≤ 2 秒')
para('• 支持分页查询，默认每页 10 条')

heading('3.4.4 订单管理模块（order）', level=3)

para('【功能描述】', bold=True)
para('提供订单创建、支付、发货、确认收货的完整订单生命周期管理。')

para('【输入要求】', bold=True)
para('• 创建订单：商品 ID、收货人姓名、手机号、收货地址')
para('• 发货：快递公司编码、运单号')

para('【处理要求】', bold=True)
para('• 订单状态流转：CREATED → PAID → SHIPPED → COMPLETED')
para('• 订单不可重复支付')
para('• 卖家仅可对自己商品的订单进行发货操作')
para('• 订单金额以"分"为单位存储')
para('• 记录完整的状态变更日志（order_events 表）')
para('• 发货时自动创建物流运单（shipments 表）')

para('【输出要求】', bold=True)
para('• 订单详情包含：订单信息、物流运单信息、状态变更事件列表')

para('【性能指标】', bold=True)
para('• 订单查询响应时间：≤ 2 秒')
para('• 并发支持：≥ 2000 订单操作/小时')

heading('3.4.5 售后管理模块（aftersale）', level=3)

para('【功能描述】', bold=True)
para('提供订单完成后的售后申请与处理功能，支持仅退款和退货退款两种类型。')

para('【输入要求】', bold=True)
para('• 售后申请：订单 ID、售后类型（REFUND / RETURN_REFUND）、申请原因')

para('【处理要求】', bold=True)
para('• 售后状态流转：REQUESTED → APPROVED / REJECTED → CLOSED')
para('• 仅买家可对自己已支付/已完成的订单发起售后')
para('• 卖家可审批通过或拒绝售后申请')
para('• 售后通过后，如果是仅退款类型，订单状态变更为 CANCELLED')

para('【输出要求】', bold=True)
para('• 返回售后申请的完整信息，包括申请时间、处理时间、处理结果')

heading('3.4.6 物流追踪模块（logistics）', level=3)

para('【功能描述】', bold=True)
para('提供物流运单管理和物流轨迹查询功能。采用策略模式设计，当前使用 Mock 实现，预留真实物流 API 接入接口。')

para('【输入要求】', bold=True)
para('• 查询物流：订单 ID（从路径获取）')

para('【处理要求】', bold=True)
para('• 运单状态：CREATED → IN_TRANSIT → DELIVERED')
para('• 物流查询接口无需认证（公开查询）')
para('• 使用策略模式（LogisticsProvider 接口 → MockLogisticsProvider 实现）')
para('• 支持多家快递公司编码')

para('【输出要求】', bold=True)
para('• 返回物流轨迹列表，包含时间、地点、状态描述')

heading('3.4.7 评论互动模块（comment）', level=3)

para('【功能描述】', bold=True)
para('提供商品评论功能，买家可对商品发表文字评价。')

para('【处理要求】', bold=True)
para('• 评论与商品和用户关联')
para('• 评论在商品详情页展示')

heading('3.4.8 收藏关注模块（favorite）', level=3)

para('【功能描述】', bold=True)
para('提供商品收藏与取消收藏功能，用户可在个人中心查看收藏列表。')

para('【处理要求】', bold=True)
para('• 同一用户对同一商品不可重复收藏')
para('• 支持收藏列表分页查询')
para('• 支持一键取消收藏')

heading('3.4.9 私聊通讯模块（chat）', level=3)

para('【功能描述】', bold=True)
para('提供买卖双方一对一私聊功能，支持消息中心和会话列表管理。')

para('【处理要求】', bold=True)
para('• 按会话组织消息，每个会话对应一对买卖双方')
para('• 支持消息已读/未读状态')
para('• 支持实时消息通知（前端轮询）')

heading('3.4.10 出价议价模块（offer）', level=3)

para('【功能描述】', bold=True)
para('提供买家出价（还价）功能，卖家可选择接受或拒绝。')

para('【处理要求】', bold=True)
para('• 出价状态：PENDING → ACCEPTED / REJECTED')
para('• 买家可对在售商品提出还价价格')
para('• 卖家收到出价后可选择接受或拒绝')

heading('3.4.11 评分评价模块（rating）', level=3)

para('【功能描述】', bold=True)
para('提供交易完成后买家对卖家的评分功能，评分在卖家主页展示。')

para('【处理要求】', bold=True)
para('• 评分与订单关联，一笔订单只能评分一次')
para('• 计算卖家综合评分（平均分）')
para('• 在卖家商品列表页展示评分信息')

heading('3.4.12 举报管理模块（report）', level=3)

para('【功能描述】', bold=True)
para('提供用户举报功能和管理员举报后台处理功能。')

para('【处理要求】', bold=True)
para('• 举报包含多种原因类型（虚假信息、违禁品、欺诈行为等）')
para('• 举报状态：PENDING → HANDLED / DISMISSED')
para('• 管理员可在后台审核和处理举报')
para('• 同一用户对同一对象不可重复举报（防刷机制）')

heading('3.4.13 管理后台模块（admin）', level=3)

para('【功能描述】', bold=True)
para('提供管理员专属后台界面，实现对平台数据和业务的全面管理。')

para('【处理要求】', bold=True)
para('• 仅 ADMIN 角色用户可访问后台')
para('• 仪表盘：展示用户总数、商品总数、订单总数、待处理举报数等统计')
para('• 用户管理：查看、搜索、禁用/启用用户')
para('• 商品管理：查看、搜索、下架/上架所有商品')
para('• 订单管理：查看所有订单详情和状态')
para('• 售后管理：处理平台所有售后申请')
para('• 举报管理：审核和处理用户举报')

heading('3.4.14 支付模块（payment）', level=3)

para('【功能描述】', bold=True)
para('提供支付占位接口，当前为模拟实现。后续可接入支付宝、微信支付等真实支付渠道。')

para('【处理要求】', bold=True)
para('• 当前实现：POST /api/orders/{id}/mark-paid 直接变更订单状态为 PAID')
para('• 预留真实支付接口扩展点')

heading('3.5 CSCI 外部接口需求', level=2)

heading('3.5.1 接口标识和接口图', level=3)
para('本系统前端通过 HTTP/HTTPS 协议调用后端 RESTful API。所有 API 统一前缀为 /api。')
para('前端开发服务器（Vite）配置了代理转发，将 /api 请求自动转发至后端 http://localhost:8088，'
    '开发环境下无需处理跨域问题。')

para('系统对外提供以下 API 端点分类：')

add_table(
    ['路径前缀', '模块', '是否需认证', '说明'],
    [
        ['/api/auth/**', '用户认证', '否（登录/注册）/ 是（改密）', '注册、登录、密码修改'],
        ['/api/auth/me', '当前用户', '是', '获取当前登录用户信息'],
        ['GET /api/products', '商品查询', '否', '公开浏览商品列表'],
        ['/api/products/**', '商品管理', '是', '商品 CRUD 操作'],
        ['/api/categories/**', '商品分类', '否', '获取商品分类'],
        ['/api/orders/**', '订单管理', '是', '订单 CRUD 操作'],
        ['/api/after-sale/**', '售后管理', '是', '售后申请与处理'],
        ['GET /api/shipments/*/track', '物流查询', '否', '公开查询物流轨迹'],
        ['/api/comments/**', '商品评论', '是', '评论发布与查看'],
        ['/api/favorites/**', '收藏管理', '是', '收藏与取消收藏'],
        ['/api/chat/**', '私聊通讯', '是', '消息发送与查询'],
        ['/api/offers/**', '出价议价', '是', '买家的出价与卖家处理'],
        ['/api/ratings/**', '评分评价', '是', '交易评分'],
        ['/api/reports/**', '举报管理', '是', '提交举报'],
        ['/api/admin/**', '管理后台', '是（ADMIN 角色）', '管理员专用接口'],
        ['/api/user/**', '用户管理', '是', '个人信息与地址管理'],
        ['/api/regions/**', '地区数据', '否', '省市区查询'],
        ['/swagger', 'API 文档', '否', 'Swagger UI 接口文档'],
    ]
)

heading('3.5.2 认证接口规范', level=3)
para('认证方式：JWT Bearer Token')
para('请求头格式：Authorization: Bearer eyJhbGciOiJIUzI1NiJ9...')
para('令牌获取：通过 /api/auth/login 或 /api/auth/register 接口获得')
para('令牌有效期：120 分钟（可配置）')

heading('3.6 CSCI 内部接口需求', level=2)
para('系统后端内部模块间通过 Spring 依赖注入进行协作，模块间接口为 Java 接口（Interface）和方法调用。主要内部接口如下：')

add_table(
    ['接口', '提供方', '调用方', '说明'],
    [
        ['JwtService', 'auth.security', 'auth.service', 'JWT 令牌生成与验证'],
        ['LogisticsProvider', 'logistics.provider', 'logistics.service', '物流查询（策略模式）'],
        ['UserRepository', 'auth.repository', '各模块 Service', '用户数据查询'],
        ['ProductRepository', 'product.repository', '各模块 Service', '商品数据查询'],
        ['OrderRepository', 'order.repository', '各模块 Service', '订单数据查询'],
    ]
)

para('前端内部通过 Vue 3 的 Composition API 进行组件间通信和数据共享。全局状态使用 Pinia Store 管理（用户状态、通知状态、主题状态）。')

heading('3.7 CSCI 内部数据需求', level=2)

para('系统数据库包含以下核心数据表：')

add_table(
    ['表名', '对应实体', '主要字段', '说明'],
    [
        ['users', 'User', 'id, nickname, password_hash, status, role, phone, email, avatar_url', '用户基本信息表'],
        ['user_identities', 'UserIdentity', 'id, user_id, identity_type, identifier, verified', '用户登录标识表（手机/邮箱）'],
        ['user_addresses', 'UserAddress', 'id, user_id, receiver_name, phone, province/city/district_code, detail', '用户收货地址表'],
        ['products', 'Product', 'id, seller_id, title, price_cent, cover_image_url, category_id, description, quantity, status, condition, free_shipping, shipping_fee_cent', '商品信息表'],
        ['product_images', 'ProductImage', 'id, product_id, image_url, sort_order', '商品多图表'],
        ['categories', 'Category', 'id, name, parent_id, sort_order', '商品分类表（支持多级）'],
        ['orders', 'Order', 'id, buyer_id, seller_id, product_id, amount_cent, status, receiver_*, address_id, paid/shipped/completed/cancelled/settled_at', '订单信息表'],
        ['order_events', 'OrderEvent', 'id, order_id, from_status, to_status, note', '订单状态变更日志表'],
        ['shipments', 'Shipment', 'id, order_id, carrier_code, tracking_no, status', '物流运单表（1:1订单）'],
        ['after_sale_requests', 'AfterSaleRequest', 'id, order_id, request_type, reason, status, requested_at, handled_at', '售后申请表'],
        ['comments', 'Comment', 'id, product_id, user_id, content', '商品评论表'],
        ['favorites', 'Favorite', 'id, user_id, product_id', '商品收藏表'],
        ['chat_messages', 'ChatMessage', 'id, sender_id, receiver_id, product_id, content, is_read', '私聊消息表'],
        ['offers', 'Offer', 'id, product_id, buyer_id, seller_id, price_cent, status', '出价议价表'],
        ['ratings', 'Rating', 'id, order_id, reviewer_id, target_user_id, score, comment', '用户评分表'],
        ['reports', 'Report', 'id, reporter_id, target_type, target_id, reason, description, status', '举报信息表'],
    ]
)

para('数据库索引设计：')
para('• user_identities: UNIQUE(identity_type, identifier)')
para('• products: INDEX(status), INDEX(category_id)')
para('• orders: INDEX(buyer_id), INDEX(seller_id)')
para('• shipments: UNIQUE(order_id)')
para('• favorites: UNIQUE(user_id, product_id)')
para('• chat_messages: INDEX(sender_id, receiver_id)')

heading('3.8 适应性需求', level=2)
para('（1）数据库适应性：系统 JPA 实体与 MySQL 数据库表通过 ddl-auto: update 自动同步，'
    '新增字段时无需手动修改数据库结构。切换数据库类型（如 MySQL → H2）仅需修改 application.yml 中的 datasource 配置。')
para('（2）物流服务适应性：物流查询采用策略模式设计接口，可通过添加新的 LogisticsProvider 实现类切换至真实物流 API，无需修改业务层代码。')
para('（3）支付服务适应性：支付模块预留了扩展接口，可接入支付宝、微信支付等第三方支付平台。')
para('（4）前端适配性：前端采用响应式设计，适配 PC 端和移动端不同屏幕尺寸。')
para('（5）CDN 支持：配置文件预留了 CDN 基础 URL 配置项（app.cdn-base-url），可将静态资源迁移至 CDN。')

heading('3.9 安全性需求', level=2)
para('（1）身份认证：所有敏感 API 均需携带有效 JWT 令牌，令牌过期后需重新登录获取。')
para('（2）密码安全：用户密码使用 BCrypt 算法进行哈希加密存储，系统不保存明文密码。')
para('（3）访问控制：区分普通用户（USER）和管理员（ADMIN）角色，管理员后台接口需 ADMIN 角色权限。')
para('（4）数据隔离：用户仅可操作自己的资源（修改自己发布的商品、查看自己的订单等），通过 Service 层进行身份校验。')
para('（5）接口防护：登录接口配置了速率限制（Rate Limit），防止暴力破解攻击。')
para('（6）JWT 安全：使用 HMAC-SHA256 签名算法，密钥通过配置文件注入（生产环境需替换为强密钥）。')
para('（7）SQL 注入防护：使用 Spring Data JPA 的预编译查询，避免 SQL 注入风险。')
para('（8）XSS 防护：前端 Vue 框架默认对输出进行 HTML 转义，防止跨站脚本攻击。')

heading('3.10 保密性和隐私需求', level=2)
para('（1）用户密码以 BCrypt 哈希形式存储，且通过 @JsonIgnore 注解禁止序列化至 JSON 响应。')
para('（2）用户个人信息（手机号、邮箱、地址）仅对用户本人和必要场景（如订单中的收货信息对卖家可见）展示。')
para('（3）私聊消息仅发送方和接收方可查看，其他用户无权访问。')
para('（4）系统遵循《个人信息保护法》的要求，在隐私政策页面（/policy/privacy）中说明用户数据的收集、使用和保护方式。')
para('（5）用户举报信息仅举报人和管理员可见，保护举报人隐私。')

heading('3.11 CSCI 环境需求', level=2)

para('【硬件环境】', bold=True)
para('服务器：')
para('• CPU：x86_64 架构，主频 1.5 GHz 以上')
para('• 内存：最低 2 GB（推荐 4 GB 以上，取决于并发用户数）')
para('• 存储：最低 1 GB 可用空间（不含数据库数据文件）')
para('• 网络：需要稳定的 TCP/IP 网络连接')

para('客户端：')
para('• 任何能运行现代浏览器的设备（PC、平板、手机）')
para('• 内存：最低 512 MB')
para('• 屏幕分辨率：最低 375px × 667px（移动端适配）')

para('【软件环境】', bold=True)
para('服务器运行环境：')
para('• 操作系统：Windows 10/11、Linux (CentOS/Ubuntu)、macOS')
para('• Java 运行环境：JDK 17 或更高版本（当前项目使用 JDK 24）')
para('• 数据库：MySQL 8.0+ 或 9.x')
para('• 无需额外 Web 服务器（Spring Boot 内嵌 Tomcat）')

para('客户端运行环境：')
para('• 浏览器：Chrome 90+、Edge 90+、Firefox 88+、Safari 14+（支持 ES2020+）')
para('• 需要启用 JavaScript 和 Cookies')
para('• 网络：需要与服务器建立 TCP/IP 连接')

heading('3.12 计算机资源需求', level=2)

heading('3.12.1 服务器硬件需求', level=3)
add_table(
    ['资源', '最低配置', '推荐配置'],
    [
        ['CPU', '双核 1.5 GHz', '四核 2.0 GHz 以上'],
        ['内存', '2 GB', '4 GB 以上'],
        ['磁盘空间', '1 GB（不含数据库）', '5 GB 以上（含数据库增长）'],
        ['网络', '10 Mbps', '100 Mbps 以上'],
    ]
)

heading('3.12.2 服务器软件资源需求', level=3)
add_table(
    ['软件', '版本', '说明'],
    [
        ['操作系统', 'Windows 10/11 或 Linux', '开发环境 Windows，生产推荐 Linux'],
        ['JDK', '17 及以上', '支持 LTS 版本和最新版本'],
        ['MySQL', '8.0 及以上', '关系型数据库'],
        ['Maven', '3.9+', 'Java 项目构建（使用 Wrapper 无需手动安装）'],
        ['Node.js', '18 及以上', '前端构建（开发阶段需要）'],
    ]
)

heading('3.12.3 网络通信需求', level=3)
para('• 前端（:5173）与后端（:8088）之间通过 HTTP 协议通信')
para('• 后端（:8088）与数据库（:3306）之间通过 JDBC/MySQL 协议通信')
para('• 建议前后端部署在同一网络环境，减少网络延迟')
para('• 如需公网访问，建议使用 Nginx 反向代理配置 HTTPS')

heading('3.13 软件质量因素', level=2)

add_table(
    ['质量属性', '要求说明'],
    [
        ['功能性', '系统应实现需求规格中定义的全部功能，各模块功能完整、正确。'],
        ['可靠性', '系统应能输出正确一致的结果。在异常情况下应具有恢复能力，保证系统稳定运行。系统正常运行时间 ≥ 99.9%。'],
        ['可维护性', '代码结构应清晰模块化，遵循 Package-by-Feature 架构模式，便于后期维护和功能扩展。每个模块内部 Controller → Service → Repository 三层分离，职责明确。'],
        ['可用性', '系统应能在需要时被正常访问和操作，不会因单用户操作不当导致系统不可用。'],
        ['灵活性', '系统应具备一定灵活性，能适应需求的小幅变化。如物流模块使用策略模式，可灵活切换不同物流服务提供商。'],
        ['可移植性', '系统基于 JVM 构建，可在 Windows、Linux、macOS 等操作系统上运行。前端基于 Web 标准，可运行在任何支持现代浏览器的设备上。'],
        ['可复用性', '代码组件设计为可复用的。前端组件（Toast、Skeleton、ImageUploader、AddressPicker 等）可在不同页面引用。后端 Service 方法可被多个 Controller 调用。'],
        ['可测试性', '系统支持单元测试、集成测试和端到端测试。已配置 JUnit 5 测试框架和前端 Vitest 测试工具。'],
        ['易用性', '用户界面简洁友好，操作流程符合用户直觉。提供明确的错误提示和操作反馈。支持响应式布局，适配多终端。'],
        ['安全性', '实现 JWT 认证 + BCrypt 密码加密 + 基于角色的访问控制 + API 速率限制。'],
    ]
)

heading('3.14 设计和实现约束', level=2)
para('（1）后端框架约束：必须使用 Spring Boot 框架，版本 3.3.x。')
para('（2）前端框架约束：必须使用 Vue 3 框架，构建工具使用 Vite。')
para('（3）编程语言约束：后端使用 Java 17+，前端使用 JavaScript（ES2020+）。')
para('（4）数据库约束：使用 MySQL 关系型数据库，通过 Spring Data JPA 进行 ORM 映射。')
para('（5）认证方式约束：使用 JWT 无状态令牌进行身份认证，不使用 Session。')
para('（6）项目源码需托管于 Git 仓库，使用 Git 进行版本控制。')
para('（7）开发过程中需编写关键模块的单元测试。')
para('（8）API 文档使用 Swagger (SpringDoc OpenAPI) 自动生成。')

heading('3.15 需求', level=2)
para('（已在上文 3.4 节 CSCI 需求中详细定义，此处不再赘述。）')

heading('3.16 测试', level=2)
para('系统测试要求：')
para('（1）单元测试：对核心 Service 层业务逻辑编写 JUnit 5 单元测试，使用 Mockito 模拟依赖。')
para('（2）集成测试：对 Controller 层 API 端点编写 Spring Boot Test 集成测试。')
para('（3）前端测试：使用 Vitest 对前端组件和工具函数编写单元测试。')
para('（4）功能测试：按照需求规格逐项验证各功能模块是否按预期工作。')
para('（5）安全测试：验证未登录用户无法访问受保护接口，普通用户无法访问管理员接口。')
para('（6）边界测试：测试异常输入、空数据、超长字符串等边界条件。')

heading('3.17 交付准备', level=2)
para('系统交付物包括：')
para('（1）源代码：完整的项目源代码（Git 仓库）')
para('（2）启动脚本：一键启动脚本 start.ps1 和停止脚本 stop.ps1')
para('（3）配置文件：application.yml（含完整配置项和注释）')
para('（4）项目文档：README.md（项目说明）、guide.md（启动指南）、需求规格说明书（本文档）')
para('（5）API 文档：Swagger 在线文档 + api-reference.md')
para('（6）架构文档：architecture.md、auth-design.md、database-schema.md、implementation-guide.md')
para('（7）测试报告：售后功能测试报告.md')

heading('3.18 算法说明', level=2)
para('（1）JWT 令牌生成算法：使用 HMAC-SHA256 算法对用户 ID、角色、过期时间等 Claims 进行签名，'
    '生成格式为 header.payload.signature 的 JWT 令牌。')
para('（2）BCrypt 密码哈希算法：使用 BCrypt 哈希算法对用户密码进行加密存储，'
    '盐值（salt）自动生成并嵌入哈希结果。验证时从哈希值中提取盐值重新计算比对。')
para('（3）商品分页查询算法：使用 Spring Data JPA 的 Pageable 分页机制，'
    '将查询参数 (page, size, sort) 转换为 SQL LIMIT/OFFSET 子句，返回 Page<T> 分页对象。')
para('（4）策略模式物流查询：定义 LogisticsProvider 接口（含 query 方法），'
    'MockLogisticsProvider 返回模拟物流轨迹数据。切换真实物流 API 时创建新的 Provider 实现类即可，'
    '无需修改调用方代码。')

heading('3.19 相关人员', level=2)
add_table(
    ['角色', '人员/组织', '职责'],
    [
        ['项目负责人', '—', '整体架构设计、技术选型、进度管理'],
        ['后端开发', '—', 'Spring Boot 后端 API 开发、数据库设计'],
        ['前端开发', '—', 'Vue 3 前端页面开发、组件开发'],
        ['测试人员', '—', '编写和执行测试用例、功能验证'],
        ['指导教师', '—', '需求评审、技术指导、项目验收'],
    ]
)

heading('3.20 培训需求', level=2)
para('（1）开发人员：需熟悉 Spring Boot 3.x、Vue 3、MySQL 相关技术栈。')
para('（2）系统管理员：需了解系统启动/停止操作、MySQL 数据库基本管理、日志查看与问题排查。')
para('（3）普通用户：系统界面直观简洁，无需专门培训。可通过平台上的引导提示快速上手。')

heading('3.21 后勤保障', level=2)
para('（1）开发工具：IntelliJ IDEA（后端）、VS Code（前端）、MySQL Workbench（数据库管理）、Git（版本控制）。')
para('（2）运行环境：需确保 JDK 17+、MySQL 8.0+、Node.js 18+ 正确安装和配置。')
para('（3）维护支持：项目源码托管于 Git 仓库，便于问题追踪和版本回溯。')

heading('3.22 包装需求', level=2)
para('本项目为 Web 应用，交付形式为源代码压缩包或 Git 仓库地址。不涉及物理包装。')

heading('3.23 安装需求', level=2)
para('（1）后端安装：解压项目 → 配置 MySQL 数据库 → 修改 application.yml 中的数据库密码 → 运行 .\\mvnw.cmd spring-boot:run')
para('（2）前端安装：cd frontend → npm install → npm run dev')
para('（3）或使用一键启动脚本：.\\start.ps1')
para('（4）详细安装步骤见 guide.md 启动指南。')

heading('3.24 优先级和关键程度', level=2)

add_table(
    ['优先级', '模块', '说明'],
    [
        ['P0 关键', '用户认证模块', '系统安全基础，用户无法登录则所有功能不可用'],
        ['P0 关键', '商品管理模块', '核心业务，买卖双方交易的基础'],
        ['P0 关键', '订单管理模块', '核心业务，交易闭环的关键环节'],
        ['P0 关键', '管理后台模块', '平台运营管理必需'],
        ['P1 重要', '售后管理模块', '保障用户权益，完善交易闭环'],
        ['P1 重要', '物流追踪模块', '提升用户体验，跟踪物流状态'],
        ['P1 重要', '收藏关注模块', '提升用户粘性和使用体验'],
        ['P1 重要', '评论互动模块', '增加社区互动和商品信任度'],
        ['P2 增强', '私聊通讯模块', '提升买卖双方沟通效率'],
        ['P2 增强', '出价议价模块', '增加交易灵活性'],
        ['P2 增强', '评分评价模块', '建设信用评价体系'],
        ['P2 增强', '举报管理模块', '维护平台秩序和内容质量'],
    ]
)

doc.add_page_break()

# ============================================================
# 第 4 章：合格性规定
# ============================================================
heading('4 合格性规定', level=1)
para('（1）系统必须能成功启动，后端在 8088 端口正常监听 HTTP 请求，前端在 5173 端口正常提供页面服务。')
para('（2）用户可通过浏览器访问 http://localhost:5173 正常浏览首页和商品列表。')
para('（3）用户可完成注册 → 登录 → 浏览商品 → 下单 → 支付 → 查看物流 → 确认收货的完整交易流程。')
para('（4）所有 API 接口返回格式符合统一响应规范 { success, data, error }。')
para('（5）未登录用户访问需认证接口时返回 401 状态码。')
para('（6）普通用户访问管理员接口时返回 403 状态码。')
para('（7）管理员可通过 /admin 路径访问后台管理界面。')
para('（8）系统在正常负载下各接口响应时间满足 3.4 节规定的性能指标。')
para('（9）用户密码以 BCrypt 哈希形式存储，不得以明文形式出现在数据库或日志中。')

doc.add_page_break()

# ============================================================
# 第 5 章：需求可追踪性
# ============================================================
heading('5 需求可追踪性', level=1)

para('下表展示了各功能需求与实现模块、前端页面之间的追踪关系：')

add_table(
    ['需求编号', '需求名称', '后端模块/包', '前端页面/组件', '状态'],
    [
        ['REQ-001', '用户注册', 'auth/AuthController', 'views/Login.vue', '✓ 已实现'],
        ['REQ-002', '用户登录', 'auth/AuthController', 'views/Login.vue', '✓ 已实现'],
        ['REQ-003', 'JWT 认证', 'auth/security/*', 'api.js (拦截器)', '✓ 已实现'],
        ['REQ-004', '密码修改', 'auth/AuthController', '—', '✓ 已实现'],
        ['REQ-005', '角色访问控制', 'auth/security/*', 'router.js (路由守卫)', '✓ 已实现'],
        ['REQ-006', '个人资料', 'user/UserController', 'views/Profile.vue', '✓ 已实现'],
        ['REQ-007', '地址管理', 'user/AddressController', 'views/AddressForm.vue', '✓ 已实现'],
        ['REQ-008', '地区选择', 'user/RegionController', 'components/RegionCascader.vue', '✓ 已实现'],
        ['REQ-009', '商品发布', 'product/ProductController', 'views/Sell.vue', '✓ 已实现'],
        ['REQ-010', '商品列表', 'product/ProductController', 'views/Home.vue', '✓ 已实现'],
        ['REQ-011', '商品详情', 'product/ProductController', 'views/ProductDetail.vue', '✓ 已实现'],
        ['REQ-012', '商品编辑', 'product/ProductController', 'views/Sell.vue', '✓ 已实现'],
        ['REQ-013', '商品分类', 'product/category/*', 'components/CategoryNav.vue', '✓ 已实现'],
        ['REQ-014', '商品多图', 'product/image/*', 'components/ImageUploader.vue, ImageGallery.vue', '✓ 已实现'],
        ['REQ-015', '商品收藏', 'favorite/FavoriteController', 'views/MyFavorites.vue', '✓ 已实现'],
        ['REQ-016', '订单创建', 'order/OrderController', 'views/ProductDetail.vue', '✓ 已实现'],
        ['REQ-017', '订单支付', 'order/OrderController, payment/*', 'views/Order.vue', '✓ 已实现'],
        ['REQ-018', '订单发货', 'order/OrderController', 'views/MyOrders.vue', '✓ 已实现'],
        ['REQ-019', '订单列表', 'order/OrderController', 'views/MyOrders.vue', '✓ 已实现'],
        ['REQ-020', '订单详情', 'order/OrderController', 'views/Order.vue', '✓ 已实现'],
        ['REQ-021', '售后申请', 'aftersale/AfterSaleController', 'views/MyAfterSales.vue', '✓ 已实现'],
        ['REQ-022', '售后处理', 'aftersale/AfterSaleController', '—', '✓ 已实现'],
        ['REQ-023', '物流查询', 'logistics/LogisticsController', 'views/Order.vue', '✓ 已实现'],
        ['REQ-024', '商品评论', 'comment/CommentController', 'views/ProductDetail.vue', '✓ 已实现'],
        ['REQ-025', '私聊通讯', 'chat/ChatMessageController', 'views/Messages.vue', '✓ 已实现'],
        ['REQ-026', '出价议价', 'offer/OfferController', 'views/ProductDetail.vue', '✓ 已实现'],
        ['REQ-027', '用户评分', 'rating/RatingController', 'views/ProductDetail.vue', '✓ 已实现'],
        ['REQ-028', '举报提交', 'report/ReportController', 'views/ProductDetail.vue', '✓ 已实现'],
        ['REQ-029', '卖家商品', 'product/MyProductController', 'views/SellerProducts.vue', '✓ 已实现'],
        ['REQ-030', '我的商品', 'product/MyProductController', 'views/MyProducts.vue', '✓ 已实现'],
        ['REQ-031', '管理仪表盘', 'admin/AdminDashboardController', 'views/admin/AdminDashboard.vue', '✓ 已实现'],
        ['REQ-032', '用户管理(管理员)', 'admin/AdminUserController', 'views/admin/AdminUsers.vue', '✓ 已实现'],
        ['REQ-033', '商品管理(管理员)', 'admin/AdminProductController', 'views/admin/AdminProducts.vue', '✓ 已实现'],
        ['REQ-034', '订单管理(管理员)', 'admin/AdminOrderController', 'views/admin/AdminOrders.vue', '✓ 已实现'],
        ['REQ-035', '售后管理(管理员)', 'admin/AdminAfterSaleController', 'views/admin/AdminAfterSale.vue', '✓ 已实现'],
        ['REQ-036', '举报管理(管理员)', 'admin/AdminReportController', 'views/admin/AdminReports.vue', '✓ 已实现'],
        ['REQ-037', '售后政策页面', '—', 'views/AfterSalePolicy.vue', '✓ 已实现'],
        ['REQ-038', '隐私政策页面', '—', 'views/PrivacyPolicy.vue', '✓ 已实现'],
        ['REQ-039', '切换账号', '—', 'views/SwitchAccount.vue', '✓ 已实现'],
    ]
)

doc.add_page_break()

# ============================================================
# 第 6 章：待解决问题
# ============================================================
heading('6 待解决问题', level=1)
para('（1）真实支付集成：当前支付功能为模拟实现（mark-paid 占位），后续需对接支付宝/微信支付等真实支付渠道。')
para('（2）真实物流 API：当前物流查询使用 Mock 模拟数据，后续需接入快递 100 或菜鸟物流等真实物流查询 API。')
para('（3）数据库迁移管理：当前使用 ddl-auto: update 自动同步表结构，不适合生产环境。'
    '后续应引入 Flyway 或 Liquibase 进行版本化的数据库迁移管理。')
para('（4）消息推送实时化：当前私聊消息采用轮询方式获取，后续可引入 WebSocket 实现实时消息推送。')
para('（5）搜索引擎优化：当前商品搜索基于数据库 LIKE 查询，后续可引入 Elasticsearch 实现全文搜索引擎。')
para('（6）缓存优化：当前未使用缓存，后续可引入 Redis 缓存热门商品列表和会话数据。')
para('（7）容器化部署：后续可编写 Dockerfile 和 docker-compose.yml，实现容器化一键部署。')
para('（8）CDN 集成：配置文件已预留 CDN 基础 URL 配置，后续可将图片等静态资源迁移至 CDN。')

doc.add_page_break()

# ============================================================
# 第 7 章：注解
# ============================================================
heading('7 注解', level=1)

para('7.1 缩略语说明')
para('• SRS：Software Requirements Specification，软件需求规格说明书')
para('• CSCI：Computer Software Configuration Item，计算机软件配置项')
para('• API：Application Programming Interface，应用程序编程接口')
para('• JWT：JSON Web Token，JSON Web 令牌')
para('• SPA：Single Page Application，单页面应用程序')
para('• ORM：Object-Relational Mapping，对象关系映射')
para('• JPA：Java Persistence API，Java 持久化 API')
para('• JDBC：Java Database Connectivity，Java 数据库连接')
para('• CRUD：Create, Read, Update, Delete，增删改查')
para('• REST：Representational State Transfer，表述性状态转移')

para('7.2 文档约定')
para('• 本文档中的"用户"均指平台的最终使用者，包括买家和卖家。')
para('• 本文档中的"管理员"指具有 ADMIN 角色的后台管理用户。')
para('• API 路径中的 {id} 表示路径变量，实际请求时替换为具体的资源 ID。')
para('• 金额单位"分"（cent）是为了避免浮点数精度问题，前端展示时自动转换为元。')

para('7.3 参考项目结构')
para('完整的项目源码目录结构如下：')
para(
    '├── backend/                          # Spring Boot 后端\n'
    '│   └── src/main/java/com/secondhand/\n'
    '│       ├── App.java                  # 启动入口\n'
    '│       ├── admin/                    # 管理后台模块\n'
    '│       ├── aftersale/                # 售后模块\n'
    '│       ├── auth/                     # 认证模块\n'
    '│       ├── chat/                     # 私聊模块\n'
    '│       ├── comment/                  # 评论模块\n'
    '│       ├── common/                   # 公共基础设施\n'
    '│       ├── config/                   # 全局配置\n'
    '│       ├── favorite/                 # 收藏模块\n'
    '│       ├── logistics/                # 物流模块\n'
    '│       ├── offer/                    # 出价模块\n'
    '│       ├── order/                    # 订单模块\n'
    '│       ├── payment/                  # 支付模块\n'
    '│       ├── product/                  # 商品模块\n'
    '│       │   ├── category/             # 商品分类\n'
    '│       │   └── image/                # 商品图片\n'
    '│       ├── rating/                   # 评分模块\n'
    '│       ├── report/                   # 举报模块\n'
    '│       └── user/                     # 用户模块\n'
    '└── frontend/                         # Vue 3 前端\n'
    '    └── src/\n'
    '        ├── App.vue                   # 主布局\n'
    '        ├── api.js                    # HTTP 请求封装\n'
    '        ├── router.js                 # 路由配置\n'
    '        ├── components/               # 通用组件\n'
    '        ├── stores/                   # Pinia 状态管理\n'
    '        └── views/                    # 页面组件\n'
    '            └── admin/                # 管理后台页面\n'
)

# ===== 保存 =====
output_path = r'D:\Second-hand Transaction\需求规格报告.docx'
doc.save(output_path)
print(f'文档已保存至: {output_path}')
print('Done!')
